"""Text processing functions"""
import io
import json
import os
import re
import string
import urllib
from typing import Dict, Generator, Optional

import markdown
from docx import Document
from html2docx import html2docx
from md2pdf.core import md2pdf
from selenium.webdriver.remote.webdriver import WebDriver
from weasyprint import HTML

from app.config import Config as GlobalConfig
from app.utils.production import Production

from ..actions.tables import add_table_to_doc, tables_to_html
from ..agent.llm_utils import create_chat_completion
from ..config import Config

CFG = Config()


def split_text(text: str, max_length: int = 8192) -> Generator[str, None, None]:
    """Split text into chunks of a maximum length

    Args:
        text (str): The text to split
        max_length (int, optional): The maximum length of each chunk. Defaults to 8192.

    Yields:
        str: The next chunk of text

    Raises:
        ValueError: If the text is longer than the maximum length
    """
    paragraphs = text.split("\n")
    current_length = 0
    current_chunk = []

    for paragraph in paragraphs:
        if current_length + len(paragraph) + 1 <= max_length:
            current_chunk.append(paragraph)
            current_length += len(paragraph) + 1
        else:
            yield "\n".join(current_chunk)
            current_chunk = [paragraph]
            current_length = len(paragraph) + 1

    if current_chunk:
        yield "\n".join(current_chunk)


def summarize_text(
    url: str, text: str, question: str, driver: Optional[WebDriver] = None
) -> str:
    """Summarize text using the OpenAI API

    Args:
        url (str): The url of the text
        text (str): The text to summarize
        question (str): The question to ask the model
        driver (WebDriver): The webdriver to use to scroll the page

    Returns:
        str: The summary of the text
    """
    if not text:
        return "Error: No text to summarize"

    summaries = []
    chunks = list(split_text(text))
    scroll_ratio = 1 / len(chunks)

    print(f"Summarizing url: {url} with total chunks: {len(chunks)}")
    for i, chunk in enumerate(chunks):
        if driver:
            scroll_to_percentage(driver, scroll_ratio * i)

        # memory_to_add = f"Source: {url}\n" f"Raw content part#{i + 1}: {chunk}"

        # MEMORY.add_documents([Document(page_content=memory_to_add)])

        messages = [create_message(chunk, question)]

        summary = create_chat_completion(
            model=CFG.fast_llm_model,
            messages=messages,
            max_tokens=CFG.summary_token_limit,
        )
        summaries.append(summary)
        # memory_to_add = f"Source: {url}\n" f"Content summary part#{i + 1}: {summary}"

        # MEMORY.add_documents([Document(page_content=memory_to_add)])

    combined_summary = "\n".join(summaries)
    messages = [create_message(combined_summary, question)]

    final_summary = create_chat_completion(
        model=CFG.fast_llm_model, messages=messages, max_tokens=CFG.summary_token_limit
    )
    print("Final summary length: ", len(combined_summary))
    print(final_summary)

    return final_summary


def scroll_to_percentage(driver: WebDriver, ratio: float) -> None:
    """Scroll to a percentage of the page

    Args:
        driver (WebDriver): The webdriver to use
        ratio (float): The percentage to scroll to

    Raises:
        ValueError: If the ratio is not between 0 and 1
    """
    if ratio < 0 or ratio > 1:
        raise ValueError("Percentage should be between 0 and 1")
    driver.execute_script(f"window.scrollTo(0, document.body.scrollHeight * {ratio});")


def create_message(chunk: str, question: str) -> Dict[str, str]:
    """Create a message for the chat completion

    Args:
        chunk (str): The chunk of text to summarize
        question (str): The question to answer

    Returns:
        Dict[str, str]: The message to send to the chat completion
    """
    return {
        "role": "user",
        "content": f'"""{chunk}""" Using the above text, answer in short the following'
        f' question: "{question}" -- if the question cannot be answered using the text,'
        " simply summarize the text. "
        "Include all factual information, numbers, stats etc if available.",
    }


def write_to_file(filename: str, text: str) -> None:
    """Write text to a file

    Args:
        text (str): The text to write
        filename (str): The filename to write to
    """
    with open(filename, "w", encoding="cp437", errors="ignore") as file:
        file.write(text)


async def save_markdown(task: str, path: str, text: str) -> str:
    """
    The function `save_markdown` saves a markdown file to a specified path, either in a production or
    development environment.

    Args:
      task (str): A string representing the task or purpose of the markdown file.
      path (str): The `path` parameter is a string that represents the file path where the markdown file
    will be saved.
      text (str): The `text` parameter is a string that represents the content of the markdown file that
    you want to save.

    Returns:
      the encoded file path.
    """
    if GlobalConfig.GCP_PROD_ENV:
        encoded_file_path = await _save_markdown_prod(task, path, text)

    else:
        encoded_file_path = await _save_markdown_dev(task, path, text)

    return encoded_file_path


async def _save_markdown_prod(task: str, path: str, text: str) -> str:
    """
    The function `_save_markdown_prod` saves a markdown file to a specified path in a user's production
    bucket and returns the encoded file path.

    Args:
      task (str): The `task` parameter is a string that represents the name or identifier of the task.
    It is used to create the file path for the markdown file.
      path (str): The `path` parameter is a string that represents the directory path where the markdown
    file will be saved.
      text (str): The `text` parameter is a string that represents the content of the markdown file that
    you want to save.

    Returns:
      the encoded file path of the uploaded markdown file.
    """
    file_path = f"{path}/{task}"
    user_bucket = Production.get_users_bucket()
    blob = user_bucket.blob(f"{file_path}.md")
    blob.upload_from_string(text, content_type="text/markdown")

    encoded_file_path = urllib.parse.quote(f"{file_path}.md")

    return encoded_file_path


async def _save_markdown_dev(task: str, path: str, text: str) -> str:
    """
    The function `_save_markdown_dev` saves a markdown text to a specified file path and returns the
    encoded file path.

    Args:
      task (str): The `task` parameter is a string that represents the name or identifier of the task or
    report. It is used to create the file name for the markdown file.
      path (str): The `path` parameter is a string that represents the directory path where the markdown
    file will be saved.
      text (str): The `text` parameter is a string that represents the content of the markdown file that
    you want to save.

    Returns:
      the encoded file path of the saved markdown file.
    """
    # Ensure that this file path exists
    os.makedirs(path, exist_ok=True)

    # Get the complete file path based reports folder, type of report
    file_path = os.path.join(path, task)

    # Write the report to markdown file
    write_to_file(f"{file_path}.md", text)

    encoded_file_path = urllib.parse.quote(f"{file_path}.md")

    return encoded_file_path


async def write_md_to_word(task: str, path: str, report: str, tables: list):
    """
    The function `write_md_to_word` writes markdown text to a Word document and returns the encoded file
    path.

    Args:
      task (str): A string representing the task or purpose of the document.
      path (str): The `path` parameter is a string that represents the file path where the Word document
    will be saved.
      report (str): The `report` parameter is a string that represents the content of the Markdown file that
    you want to convert to a Word document.

    Returns:
      the encoded file path.
    """
    if GlobalConfig.GCP_PROD_ENV:
        encoded_file_path = await _write_md_to_word_prod(task, path, report, tables)
    else:
        encoded_file_path = await _write_md_to_word_dev(task, path, report, tables)

    return encoded_file_path


async def _write_md_to_word_prod(
    task: str, path: str, report: str, tables: list
) -> str:
    file_path = f"{path}/{task}"
    user_bucket = Production.get_users_bucket()

    # Convert Markdown to HTML
    html = markdown.markdown(report)

    # html2docx() returns an io.BytesIO() object. The HTML must be valid.
    doc_bytes = html2docx(html, title=f"{task}.docx")
    # print("doc_bytes : ", doc_bytes)
    # blob = user_bucket.blob(f"{file_path}.docx")
    # blob.upload_from_string(doc_bytes.getvalue())

    doc = Document(doc_bytes)

    if len(tables):
        # Add "Data Tables" as a title before the tables section
        doc.add_heading("Data Tables", level=1)

        # Adding tables to the Word document
        for tables_set in tables:
            tables_in_url = tables_set["tables"]
            url = tables_set["url"]
            for table_data in tables_in_url:
                print("Table title : ", table_data["title"])
                add_table_to_doc(doc, table_data["title"], table_data["values"], url)
                doc.add_paragraph()  # Adding an extra paragraph between tables

    # Create a temporary file-like object to save the updated document
    temp_doc_io = io.BytesIO()
    doc.save(temp_doc_io)
    temp_doc_io.seek(0)  # Reset the pointer to the beginning of the stream

    # Upload the Docx file to the bucket
    blob = user_bucket.blob(f"{file_path}.docx")
    blob.upload_from_file(temp_doc_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    print(f"🎉 {task} written to {file_path}.docx")

    encoded_file_path = urllib.parse.quote(f"{file_path}.docx")

    return encoded_file_path


async def _write_md_to_word_dev(task: str, path: str, report: str, tables: list) -> str:
    """
    The function `_write_md_to_word_dev` takes a task, path, report, and tables as input, converts the
    markdown report to HTML, converts the HTML to a Word document, adds tables to the Word document, and
    saves the document to the specified path. It then returns the encoded file path.

    Args:
      task (str): The name of the task or report.
      path (str): The `path` parameter is a string that represents the directory path where the Word
    document will be saved.
      report (str): The `report` parameter is a string that contains the markdown content that needs to
    be converted to a Word document.
      tables (list): The `tables` parameter is a list of dictionaries. Each dictionary represents a set
    of tables related to a specific URL. Each dictionary has two keys:

    Returns:
      the encoded file path of the generated Word document.
    """
    # Ensure that this file path exists
    os.makedirs(path, exist_ok=True)

    # Get the complete file path based reports folder, type of report
    file_path = os.path.join(path, task)

    html = markdown.markdown(report)

    # html2docx() returns an io.BytesIO() object. The HTML must be valid.
    buf = html2docx(html, title=f"{task}.docx")

    # with open(f"{file_path}.docx", "wb") as fp:
    #     fp.write(buf.getvalue())

    doc = Document(buf)
    # Add "Data Tables" as a title before the tables section
    doc.add_heading("Data Tables", level=1)

    if len(tables):
        # Adding tables to the Word document
        for tables_set in tables:
            tables_in_url = tables_set["tables"]
            url = tables_set["url"]
            for table_data in tables_in_url:
                print("Table title : ", table_data["title"])
                add_table_to_doc(doc, table_data["title"], table_data["values"], url)
                doc.add_paragraph()  # Adding an extra paragraph between tables

    doc.save(f"{file_path}.docx")

    print(f"🎉 {task} written to {file_path}.docx")

    encoded_file_path = urllib.parse.quote(f"{file_path}.docx")

    return encoded_file_path


async def write_md_to_pdf(task: str, path: str, report: str, tables: list) -> str:
    """
    The function `write_md_to_pdf` writes markdown text to a PDF file and returns the encoded file path.

    Args:
      task (str): A string representing the task or purpose of the PDF file.
      path (str): The `path` parameter is a string that represents the file path where the PDF file will
    be saved.
      report (str): The `report` parameter is a string that represents the content of the Markdown file that
    you want to convert to a PDF.

    Returns:
      the encoded file path as a string.
    """
    if GlobalConfig.GCP_PROD_ENV:
        encoded_file_path = await _write_md_to_pdf_prod(task, path, report, tables)

    else:
        encoded_file_path = await _write_md_to_pdf_dev(task, path, report, tables)

    return encoded_file_path


async def _write_md_to_pdf_prod(task: str, path: str, report: str, tables: list) -> str:
    """
    The function `_write_md_to_pdf_prod` takes a task, path, and text as input, uploads the text as a
    Markdown file to a user's bucket, converts the Markdown to HTML, generates a PDF file from the HTML
    using WeasyPrint, uploads the PDF file to the user's bucket, and returns the encoded file path of
    the PDF file.

    Args:
      task (str): The `task` parameter is a string that represents the name or identifier of the task
    for which the Markdown file and PDF file will be created.
      path (str): The `path` parameter is a string representing the directory path where the files will
    be stored.
      report (str): The `report` parameter is a string that contains the content of the Markdown file that
    needs to be converted to PDF.

    Returns:
      the encoded file path of the PDF file that was generated.
    """
    file_path = f"{path}/{task}"
    user_bucket = Production.get_users_bucket()

    # Combined html
    html = get_combined_html(report, tables)

    # Create a WeasyPrint HTML object
    html_obj = HTML(string=html)

    # Generate the PDF file
    pdf_bytes = html_obj.write_pdf()
    blob = user_bucket.blob(f"{file_path}.pdf")
    blob.upload_from_string(pdf_bytes, content_type="application/pdf")

    print(f"🎉 {task} written to {file_path}.pdf")

    encoded_file_path = urllib.parse.quote(f"{file_path}.pdf")

    return encoded_file_path


async def _write_md_to_pdf_dev(task: str, path: str, report: str, tables: list) -> str:
    """
    The function `_write_md_to_pdf_dev` takes a task, path, and text as input, writes the text to a
    markdown file, converts the markdown to HTML, generates a PDF file using WeasyPrint, and returns the
    encoded file path of the PDF.

    Args:
      task (str): The `task` parameter is a string that represents the name or identifier of the task or
    report. It is used to create the file name for the markdown and PDF files.
      path (str): The `path` parameter is the directory path where the PDF file will be saved.
      report (str): The `report` parameter is a string that contains the content of the report in Markdown
    format.

    Returns:
      the encoded file path of the generated PDF file.
    """
    # Ensure that this file path exists
    os.makedirs(path, exist_ok=True)

    # Get the complete file path based reports folder, type of report
    file_path = os.path.join(path, task)

    # print("📡 path : ", path)
    # print("📡 file_path : ", file_path)

    # Combined html
    html = get_combined_html(report, tables)

    # Create a WeasyPrint HTML object
    html_obj = HTML(string=html)

    # Generate the PDF file
    html_obj.write_pdf(f"{file_path}.pdf")

    print(f"🎉 {task} written to {file_path}.pdf")

    encoded_file_path = urllib.parse.quote(f"{file_path}.pdf")

    return encoded_file_path


def get_combined_html(report: str, tables: list):
    # Convert Markdown to HTML
    report_html = markdown.markdown(report)

    # Get the html of the tables
    tables_html = ""
    for tables_in_url in tables:
        current_tables = tables_in_url.get("tables", [])
        current_url = tables_in_url.get("url", "")
        tables_html += tables_to_html(current_tables, current_url)

    combined_html = report_html
    if len(tables_html):
        combined_html += "<br><h1>Data Tables</h1><br>" + tables_html

    return combined_html


def read_txt_files(directory, tables=False):
    """
    The function reads research text files from a specified directory, either in a production or
    development environment.

    Args:
      directory: The directory parameter is the path to the directory where the text files are located.
      tables: The `tables` parameter is a boolean flag that indicates whether or not to extract tables
    from the text files. If `tables` is set to `True`, the function will extract tables from the text
    files. If `tables` is set to `False` (default), the function will only read. Defaults to False

    Returns:
      the variable "all_text".
    """
    print("📡 Reading research text files from : ", directory)

    if GlobalConfig.GCP_PROD_ENV:
        all_text = _read_text_files_prod(directory, tables)
    else:
        all_text = _read_text_files_dev(directory, tables)

    return all_text


def _read_text_files_prod(directory, tables=False):
    """
    The function `_read_text_files_prod` reads the content of text files in a specified directory and
    returns all the text concatenated together.

    Args:
      directory: The `directory` parameter is the name of the directory in the production bucket where
    the text files are located.
      tables: The `tables` parameter is a boolean flag that determines whether to include only the
    content of the "tables.txt" file or include the content of all ".txt" files in the specified
    directory. Defaults to False

    Returns:
      a string containing the content of all the text files in the specified directory.
    """
    bucket = Production.get_users_bucket()
    blobs = bucket.list_blobs(prefix=directory)
    all_text = ""
    for filename in blobs:
        print(f"📡 Filename : {filename}")
        if tables:
            if filename == "tables.txt":
                content = filename.download_as_text()
                print(f"📡 content : {content}")
                all_text += content + "\n"
        else:
            if filename.name.endswith(".txt"):
                content = filename.download_as_text()
                print(f"📡 content : {content}")
                all_text += content + "\n"

    return all_text


def _read_text_files_dev(directory, tables=False):
    """
    The function `_read_text_files_dev` reads the contents of text files in a given directory and
    returns the concatenated text.

    Args:
      directory: The directory parameter is the path to the directory where the text files are located.
      tables: The `tables` parameter is a boolean flag that determines whether to include the contents
    of a file named "tables.txt" in the `all_text` output. If `tables` is set to `True`, the function
    will only include the contents of "tables.txt" in the output. If `. Defaults to False

    Returns:
      a string containing the contents of all the text files in the specified directory. If the `tables`
    parameter is set to `True`, it only includes the contents of the "tables.txt" file.
    """
    all_text = ""
    for filename in os.listdir(directory):
        if tables:
            if filename == "tables.txt":
                with open(
                    os.path.join(directory, filename),
                    "r",
                    errors="ignore",
                ) as file:
                    all_text += file.read() + "\n"
        else:
            if filename.endswith(".txt"):
                with open(
                    os.path.join(directory, filename),
                    "r",
                    encoding="cp437",
                    errors="ignore",
                ) as file:
                    all_text += file.read() + "\n"

    return all_text


def md_to_pdf(input_file, output_file):
    md2pdf(
        output_file,
        md_content=None,
        md_file_path=input_file,
        css_file_path=None,
        base_url=None,
    )


def remove_roman_numerals(input_string):
    """
    The function `remove_roman_numerals` takes an input string and removes any occurrences of Roman
    numerals from it.

    Args:
      input_string: The input string is the string from which you want to remove Roman numerals.

    Returns:
      a string with all Roman numerals removed.
    """
    # Define a regular expression pattern for Roman numerals
    roman_numerals_pattern = r"\b[IVXLCDM]+\b"

    # Use re.sub() to replace Roman numerals with an empty string
    result_string = re.sub(roman_numerals_pattern, "", input_string)

    return result_string
