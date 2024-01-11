import os
import re

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.config import Config as GlobalConfig
from app.utils.production import Production

from ...document import add_hyperlink
from ..utils.text import *


class TableExtractor:
    def __init__(self, dir_path: str):
        self.dir_path = dir_path
        self.tables = []
        self.tables_path = os.path.join(self.dir_path, "tables.txt")

    def read_tables(self):
        if GlobalConfig.GCP_PROD_ENV:
            return read_txt_files(self.dir_path, tables=True) or ""
        else:
            if os.path.isdir(self.dir_path):
                return read_txt_files(self.dir_path, tables=True)
            else:
                return ""

    def save_tables(self):
        # save tables
        if GlobalConfig.GCP_PROD_ENV:
            user_bucket = Production.get_users_bucket()
            blob = user_bucket.blob(self.tables_path)
            blob.upload_from_string(str(self.tables))
        else:
            os.makedirs(os.path.dirname(self.tables_path), exist_ok=True)
            write_to_file(self.tables_path, str(self.tables))

    def extract_tables(self, url: str) -> list:
        def extract_table_title(table):
            def process_table_title(title):
                # Removing Table numberings like "Table 1:", "Table1-", etc.
                cleaned_title = re.sub(r'^\s*Table\s*\d+\s*[:\-]\s*', '', title, flags=re.IGNORECASE)
                return cleaned_title.strip()

            # Extract table title
            title = ""  # Default title if no specific title found
            table_caption = table.find("caption")
            if table_caption:
                title = table_caption.get_text(strip=True)
            else:
                previous_tag = table.find_previous(["h1", "h2", "h3", "h4", "h5", "h6", "p"])
                if previous_tag:
                    title = previous_tag.get_text(strip=True)

            return process_table_title(title)

        def extract_table_data(table, title):
            table_data = []
            rows = table.find_all("tr")

            # Initialize an empty dictionary to store header row data
            header_dict = {}

            # Extract headers from thead if available
            thead = table.find("thead")
            if thead:
                header_row = thead.find("tr")
                if header_row:
                    header_cells = header_row.find_all(["th", "td"])
                    # Extracting header names
                    header_names = [cell.text.strip() for cell in header_cells if cell.text.strip()]
                    if header_names:
                        # Assuming the header row contains unique names for keys
                        for index, name in enumerate(header_names or list(rows[0].values())):
                            # Assigning each header name as a key with an empty list as its value
                            header_dict[index] = name

            # Append the dictionary containing header row data to the table_data list
            if header_dict:
                table_data.append(header_dict)

            for row in rows[1:]:
                row_data = {}
                cells = row.find_all(["td", "th"])
                valid_row = False

                for idx in range(len(header_dict.keys())):
                    try:
                        cell = cells[idx]
                        cell_text = cell.text.strip()
                    except IndexError:
                        cell_text = ""

                    row_data[str(idx)] = cell_text

                    if cell_text and cell_text not in ["NA", "n/a", "na", "-", "", "NaN"]:
                        valid_row = True

                if valid_row:
                    table_data.append(row_data)

            return {"title": title, "values": table_data}

        try:
            response = requests.get(url)
            soup = BeautifulSoup(response.content, "html.parser")

            tables = soup.find_all("table")
            extracted_tables = []

            for table in tables:
                # Check if the table contains any nested tables
                nested_tables = table.find_all("table")
                if not nested_tables:  # If no nested tables, extract the main table
                    title = extract_table_title(table)

                    # Extract table data
                    table_data = extract_table_data(table, title)

                    # Exclude tables with a single row and a single data value
                    if table_data.get('values') and (len(table_data['values']) > 1 or (len(table_data['values']) == 1 and len(table_data['values'][0]) > 1)):
                        extracted_tables.append(table_data)

            return extracted_tables

        except Exception as e:
            print(f"🚩 Exception when scraping tables : {e}") 
            return []

    def tables_to_html(self, list_of_tables: list, url: str) -> str:
        try:
            html_tables = []
            print(f"Received {len(list_of_tables)} tables...")

            for table_data in list_of_tables:
                if "tables" in table_data:
                    tables = table_data["tables"]

                for table in tables:
                    title = table.get("title", "")
                    values = table.get("values", [])

                    print(f"Title : {title}")

                    # Create the table header row
                    headers = list(values[0].values())
                    header_row = (
                        "<tr><th style='font-size: 16px; border: 1px solid black; padding: 10px; text-align: center;'>"
                        + "</th><th style='border: 1px solid black; padding: 10px; text-align: center;'>".join(
                            headers
                        )
                        + "</th></tr>"
                    )
                    html_table = [header_row]

                    # Create rows with data
                    for row in values[1:]:
                        row_values = list(row.values())
                        row_str = "<tr>"
                        for val in row_values:
                            # Check if the value is missing or empty key
                            if val == "" or val is None:
                                val = "&nbsp;"  # Replace with a non-breaking space

                            # Check if the value is numerical
                            if self.is_numerical(str(val)):
                                align_style = "text-align: right;"
                            else:
                                align_style = "text-align: left;"

                            # Add border style to always have a border around the cell
                            row_str += f"<td style='font-size: 14px; border: 1px solid black; padding: 10px; {align_style}'>{str(val)}</td>"
                        row_str += "</tr>"
                        html_table.append(row_str)

                    # Create the table HTML structure
                    table_html = (
                        "<table style='border-collapse: collapse;'><tbody>"
                        + "\n".join(html_table)
                        + "</tbody></table>"
                    )

                    # Create the title with proper heading
                    title_html = (
                        f"<p style='font-weight: bold; font-size: 18px;'>{title}</p>"
                    )

                    # Add the url of the table
                    url_html = f"<br><a href='{url}' style='float: right; font-size: 12px;'>{url}</a>"

                    # Combine title and table HTML
                    final_table_html = title_html + table_html + url_html
                    html_tables.append(final_table_html)

            return "<br><br><br>".join(html_tables)

        except Exception as e:
            print(f"🚩 Exception in converting tables to html: {e}")
            return ""
    
    def add_tables_to_doc(self, document: Document):
        if not len(self.tables):
            return

        # Add "Data Tables" as a title before the tables section
        document.add_heading("Data Tables", level=1)

        # Adding tables to the Word document
        for tables_set in self.tables:
            tables_in_url = tables_set["tables"]
            url = tables_set["url"]
            for table_data in tables_in_url:
                print("Table title : ", table_data["title"])
                self.add_table_to_doc(
                    document, table_data["title"], table_data["values"], url
                )
                document.add_paragraph()  # Adding an extra paragraph between tables

    def add_table_to_doc(
        self, document: Document, table_title: str, table_values: list, url: str
    ):
        try:
            document_clone = Document()
            document_clone.add_heading(table_title, level=2)
            table = document_clone.add_table(rows=1, cols=len(table_values[0]))
            table.style = "Table Grid"  # Applying table grid style to have borders

            # Set header rows as bold and add borders
            hdr_cells = table.rows[0].cells
            for i, key in enumerate(table_values[0].keys()):
                cell = hdr_cells[i]
                cell.text = table_values[0][key]
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].runs[0].font.bold = True
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                for border in cell._element.xpath(".//*"):
                    border_val = border.attrib.get("val")
                    if border_val is not None:
                        border.attrib.clear()
                        border.attrib["val"] = border_val
                    else:
                        border.attrib["val"] = "single"

            for row_data in table_values[1:]:
                row = table.add_row().cells
                for i, key in enumerate(row_data.keys()):
                    cell = row[i]
                    cell.text = row_data[key]
                    if self.is_numerical(row_data[key]):
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    else:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                    for border in cell._element.xpath(".//*"):
                        border_val = border.attrib.get("val")
                        if border_val is not None:
                            border.attrib.clear()
                            border.attrib["val"] = border_val
                        else:
                            border.attrib["val"] = "single"

            # Adding the table URL after the table with right alignment and as a hyperlink
            p = document_clone.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            add_hyperlink(p, url, url)

            # Append the modified content to the original document
            for element in document_clone.element.body:
                document.element.body.append(element)

        except Exception as e:
            print(f"🚩 Exception in adding tables to word document : {e}")
            # Revert changes on any exception
            pass

    def get_combined_html(self, report: str):
        # Convert Markdown to HTML
        report_html = markdown.markdown(report)

        # Get the html of the tables
        tables_html = ""
        print(f"➕ Appending {len(self.tables)} tables to report...\n")
        for tables_in_url in self.tables:
            current_tables = tables_in_url.get("tables", [])
            current_url = tables_in_url.get("url", "")
            tables_html += self.tables_to_html(current_tables, current_url)

        combined_html = report_html
        if len(tables_html):
            combined_html += "<br><h1>Data Tables</h1><br>" + tables_html

        return combined_html

    @staticmethod
    def is_numerical(value: str) -> bool:
        """
        The function `is_numerical` checks if a given value is numerical, allowing for optional commas and a
        percentage sign at the end.

        Args:
        value (str): The value parameter is a string that represents a numerical value.

        Returns:
        The function is_numerical is returning a boolean value.
        """
        numerical_pattern = re.compile(r"^-?(\d{1,3}(,\d{3})*|\d+)?(\.\d+)?%?$")
        return bool(numerical_pattern.match(str(value).replace(",", "")))
