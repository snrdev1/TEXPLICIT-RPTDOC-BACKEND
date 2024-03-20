import io
import os
import re

import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, NamedStyle
from openpyxl.utils.dataframe import dataframe_to_rows

from app.config import Config as GlobalConfig
from app.utils.common import Common
from app.utils.production import Production

from ....document import add_hyperlink
from ...utils.text import *


class TableExtractor:
    def __init__(self, dir_path: str):
        self.dir_path = dir_path
        self.tables = []
        self.tables_temp_path = os.path.join(self.dir_path, "tables.txt")
        self.tables_save_path = os.path.join(self.dir_path, "tables.xlsx")

    def read_tables(self):
        if GlobalConfig.GCP_PROD_ENV:
            return read_txt_files(self.dir_path, tables=True) or ""
        else:
            if os.path.isdir(self.dir_path):
                return read_txt_files(self.dir_path, tables=True)
            else:
                return ""

    def save_tables(self):
        # save tables
        if GlobalConfig.GCP_PROD_ENV:
            user_bucket = Production.get_users_bucket()
            blob = user_bucket.blob(self.tables_temp_path)
            blob.upload_from_string(str(self.tables))
        else:
            os.makedirs(os.path.dirname(self.tables_temp_path), exist_ok=True)
            write_to_file(self.tables_temp_path, str(self.tables))

    def extract_tables(self, url: str) -> list:
        """
        Extract tables from a given URL excluding those with hyperlinks in values and skipping dataframes with zero or one rows.

        Args:
            url (str): The URL to scrape tables from.

        Returns:
            list: List of dictionaries, each representing a table with title and values.
        """
        def extract_table_title(table) -> str:
            """
            Extract and process the title of a table.

            Args:
                table: BeautifulSoup object representing a table.

            Returns:
                str: Processed title of the table.
            """
            def process_table_title(title):
                # Removing Table numberings like "Table 1:", "Table1-", etc.
                cleaned_title = re.sub(
                    r"^\\s*Table\\s*\\d+\\s*[:\\-]\\s*", "", title, flags=re.IGNORECASE)
                return cleaned_title.strip()

            # Check for a caption tag within the table
            table_caption = table.find("caption")
            if table_caption:
                title = table_caption.get_text(strip=True)
            else:
                # If no caption, look at previous tags like h1, h2, etc.
                previous_tag = table.find_previous(
                    ["h1", "h2", "h3", "h4", "h5", "h6", "p"])
                if previous_tag:
                    title = previous_tag.get_text(strip=True)
                else:
                    title = ""

            return process_table_title(title)

        def extract_table_from_pdf() -> list:
            return []

        def extract_table_from_url() -> list:
            response = requests.get(url)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, "html.parser")
            dfs = pd.read_html(response.content)

            extracted_tables = []
            for idx, df in enumerate(dfs):
                # Skip dataframes with zero or one rows
                if len(df) > 1:
                    table_title = extract_table_title(
                        soup.find_all("table")[idx])
                    # Convert NaN values to "-"
                    df_filled = df.fillna("-")
                    extracted_tables.append(
                        {"title": table_title, "values": df_filled.to_dict(orient="records")})

            return extracted_tables

        try:
            extracted_tables = []
            if url.endswith(".pdf"):
                extracted_tables = extract_table_from_pdf()
            else:
                extracted_tables = extract_table_from_url()

            return extracted_tables

        except requests.RequestException as e:
            print(f"🚩 Request Exception when scraping tables: {e}")
            return []
        except Exception as e:
            print(f"🚩 Error in extract_tables: {e}")
            return []

    def tables_to_html(self, list_of_tables: list, url: str) -> str:
        try:
            html_tables = []
            print(f"Received {len(list_of_tables)} tables...")

            for table in list_of_tables:
                title = table.get("title", "")
                values = table.get("values", [])
                if values == []:
                    continue

                # Create the table header row
                headers = list(values[0].values())
                header_row = (
                    "<tr><th style='font-size: 16px; border: 1px solid black; padding: 10px; text-align: center;'>"
                    + "</th><th style='border: 1px solid black; padding: 10px; text-align: center;'>".join(
                        headers
                    )
                    + "</th></tr>"
                )
                html_table = [header_row]

                # Create rows with data
                for row in values[1:]:
                    row_values = list(row.values())
                    row_str = "<tr>"
                    for val in row_values:
                        # Check if the value is missing or empty key
                        if val == "" or val is None:
                            val = "&nbsp;"  # Replace with a non-breaking space

                        # Check if the value is numerical
                        if self.is_numerical(str(val)):
                            align_style = "text-align: right;"
                        else:
                            align_style = "text-align: left;"

                        # Add border style to always have a border around the cell
                        row_str += f"<td style='font-size: 14px; border: 1px solid black; padding: 10px; {align_style}'>{str(val)}</td>"
                    row_str += "</tr>"
                    html_table.append(row_str)

                # Create the table HTML structure
                table_html = (
                    "<table style='border-collapse: collapse;'><tbody>"
                    + "\n".join(html_table)
                    + "</tbody></table>"
                )

                # Create the title with proper heading
                title_html = (
                    f"<p style='font-weight: bold; font-size: 18px;'>{title}</p>"
                )

                # Add the url of the table
                url_html = f"<br><a href='{url}' style='float: right; font-size: 12px;'>source</a>"

                # Combine title and table HTML
                final_table_html = title_html + table_html + url_html
                html_tables.append(final_table_html)

            return "<br><br><br>".join(html_tables)

        except Exception as e:
            Common.exception_details("TableExtractor.tables_to_html", e)
            return ""

    def add_tables_to_doc(self, document: Document):

        def add_table_to_doc(self, document: Document, table_title: str, table_values: list, url: str):
            try:
                def clean_text(text):
                    """Remove spaces, new lines, and tabs from the text."""
                    return text.replace("\n", " ").replace("\t", " ")

                document.add_heading(table_title, level=2)

                table = document.add_table(rows=1, cols=len(table_values[0]))

                table.style = "Table Grid"  # Applying table grid style to have borders

                # Set header rows as bold and add borders
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(table_values[0].keys()):
                    cell = hdr_cells[i]
                    cell.text = clean_text(table_values[0][key])
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.paragraphs[0].runs[0].font.bold = True
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                    for border in cell._element.xpath(".//*"):
                        border_val = border.attrib.get("val")
                        if border_val is not None:
                            border.attrib.clear()
                            border.attrib["val"] = border_val
                        else:
                            border.attrib["val"] = "single"

                for row_data in table_values[1:]:
                    row = table.add_row().cells
                    for i, key in enumerate(row_data.keys()):
                        cell = row[i]
                        cell.text = clean_text(row_data[key])
                        if self.is_numerical(row_data[key]):
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        else:
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                        for border in cell._element.xpath(".//*"):
                            border_val = border.attrib.get("val")
                            if border_val is not None:
                                border.attrib.clear()
                                border.attrib["val"] = border_val
                            else:
                                border.attrib["val"] = "single"

                # Adding the table URL after the table with right alignment and as a hyperlink
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                add_hyperlink(p, "source", url, True)

            except Exception as e:
                print(f"🚩 Exception in adding tables to word document : {e}")
                # Revert changes on any exception
                pass

        if not len(self.tables):
            return

        # Add "Data Tables" as a title before the tables section
        document.add_heading("Data Tables", level=1)

        # Adding tables to the Word document
        for tables_set in self.tables:
            tables_in_url = tables_set["tables"]
            url = tables_set["url"]
            for table_data in tables_in_url:
                add_table_to_doc(
                    document, table_data["title"], table_data["values"], url)

    def get_combined_html(self, report_html: str):
        # Get the html of the tables
        tables_html = ""
        for tables_in_url in self.tables:
            current_tables = tables_in_url.get("tables", [])
            current_url = tables_in_url.get("url", "")
            tables_html += self.tables_to_html(current_tables, current_url)

        combined_html = report_html
        if len(tables_html):
            print(f"➕ Appending {len(self.tables)} tables to report...\n")
            combined_html += "<br><h2>Data Tables</h2><br>" + tables_html

        return combined_html

    async def save_tables_to_excel(self):
        """
        The `save_tables_to_excel` function in Python creates an Excel workbook with multiple sheets
        containing table data and saves it either to a production bucket or locally based on the
        environment.
        
        Returns:
          The `save_tables_to_excel` method returns the encoded file path where the tables have been
        saved. This path is returned as a result of the method execution. If an exception occurs during
        the process, the method catches the exception, logs the details, and returns an empty string.
        """

        def _get_excel_workbook():
            wb = Workbook()

            # Create a custom style for the first row in the first sheet
            first_row_style = NamedStyle(name="first_row_style")
            first_row_style.font = Font(bold=True, size=14)
            wb.add_named_style(first_row_style)

            # Insert a sheet at the beginning titled 'List of tables'
            list_sheet = wb.create_sheet(title="List of tables")
            list_sheet.append(["Table Title", "URL"])

            # Set column widths for 'List of tables' sheet
            list_sheet.column_dimensions['A'].width = 30
            list_sheet.column_dimensions['B'].width = 50

            for table_data in self.tables:
                tables = table_data["tables"]
                url = table_data["url"]

                for table in tables:
                    title = _sanitize_sheet_title(table['title'])
                    values = table['values']

                    # Create a new sheet with the title of the table
                    ws = wb.create_sheet(title=title)

                    # Convert table values to DataFrame
                    df = pd.DataFrame(values)

                    # Write DataFrame to Excel sheet
                    for row in dataframe_to_rows(df, index=False, header=True):
                        try:
                            ws.append(row)
                        except ValueError as e:
                            print(f"Error appending row {row}: {e}")
                            # Delete the sheet if there's an error appending rows
                            del wb[title]
                            print(f"Deleted sheet {title}")
                            continue

                    # Adjust column widths to fit content
                    _adjust_column_widths(ws)

                    # Add table title and URL as a clickable link to the 'List of tables' sheet
                    list_sheet.append(
                        [table['title'], '=HYPERLINK("{}", "{}")'.format(url, f"{url}")])

            # Apply custom style to the first row of the 'List of tables' sheet
            for cell in list_sheet['1:1']:
                cell.style = first_row_style

            # Apply italic and blue style to column 2 starting from row 2 in 'List of tables' sheet
            italic_blue_style = NamedStyle(name="italic_blue_style")
            italic_blue_style.font = Font(
                italic=True, color="0000FF")  # Blue color code
            # Adjust column width for URL
            list_sheet.column_dimensions['B'].width = 50
            for row in list_sheet.iter_rows(min_row=2, min_col=2, max_col=2):
                for cell in row:
                    cell.style = italic_blue_style

            # Adjust column widths for 'List of tables' sheet
            _adjust_column_widths(list_sheet)

            # Remove default sheet created by openpyxl
            wb.remove(wb["Sheet"])

            return wb

        def _sanitize_sheet_title(title: str) -> str:
            """Sanitize the sheet title by replacing invalid characters with '_' and limiting the length."""
            # Replace invalid characters with '_'
            sanitized_title = re.sub(r'[\\/*?:[\]]', '_', title)
            # Limit the length of the title to 31 characters
            return sanitized_title[:30]

        def _adjust_column_widths(ws):
            """Adjust column widths based on content length."""
            for column_cells in ws.columns:
                max_length = 0
                for cell in column_cells:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except TypeError:
                        pass
                adjusted_width = (max_length + 2) * 1.2
                ws.column_dimensions[column_cells[0]
                                     .column_letter].width = adjusted_width

        def _save_tables_to_excel_prod(workbook: Workbook):
            user_bucket = Production.get_users_bucket()

            # Create a temporary file-like object to save the updated document
            temp_doc_io = io.BytesIO()
            workbook.save(temp_doc_io)
            # Reset the pointer to the beginning of the stream
            temp_doc_io.seek(0)

            # Upload the Docx file to the bucket
            blob = user_bucket.blob(self.tables_save_path)
            blob.upload_from_file(
                temp_doc_io,
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        def _save_tables_to_excel_dev(workbook: Workbook):
            # Save the Excel file
            workbook.save(self.tables_save_path)

        try:
            if not self.tables:
                return ""
            
            workbook = _get_excel_workbook()

            if GlobalConfig.GCP_PROD_ENV:
                await _save_tables_to_excel_prod(workbook)
            else:
                await _save_tables_to_excel_dev(workbook)

            # Return url encoded path where the tables are have been saved
            encoded_file_path = urllib.parse.quote(self.tables_save_path)
            
            return encoded_file_path

        except Exception as e:
            Common.exception_details("tables.save_tables_to_excel", e)
            return ""

    @staticmethod
    def is_numerical(value: str) -> bool:
        """
        The function `is_numerical` checks if a given value is numerical, allowing for optional commas and a
        percentage sign at the end.

        Args:
        value (str): The value parameter is a string that represents a numerical value.

        Returns:
        The function is_numerical is returning a boolean value.
        """
        numerical_pattern = re.compile(
            r"^-?(\d{1,3}(,\d{3})*|\d+)?(\.\d+)?%?$")
        return bool(numerical_pattern.match(str(value).replace(",", "")))
