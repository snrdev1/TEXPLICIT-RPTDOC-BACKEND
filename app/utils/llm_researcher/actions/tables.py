import re

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from ...document import add_hyperlink


def extract_tables(url: str) -> list:
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, "html.parser")

        tables = soup.find_all("table")
        extracted_tables = []

        for table in tables:
            table_data = []
            table_title = table.find_previous(["h1", "h2", "h3", "h4", "h5", "h6", "p"])

            if table_title:
                title = table_title.text.strip()
            else:
                title = "Table"  # Default title if no specific title found

            headers = [th.text.strip() for th in table.find_all("th")]
            rows = table.find_all("tr")

            # Check if the table has headers and at least one non-header row
            if len(headers) > 0 and len(rows) > 1:
                for row in rows[1:]:
                    row_data = {}
                    cells = row.find_all(["td", "th"])
                    valid_row = False  # Flag to check if the row has valid data

                    for idx, cell in enumerate(cells):
                        cell_text = cell.text.strip()

                        # Check for NA, blank, "-", or other null values
                        if cell_text and cell_text not in ["NA", "n/a", "na", "-", "", "NaN"]:
                            row_data[str(idx)] = cell_text
                            valid_row = True  # Set the flag to True for a valid row

                    # Append the row only if it contains valid data
                    if valid_row:
                        table_data.append(row_data)

                # Append the table if it has valid rows
                if len(table_data) > 0:
                    extracted_table = {"title": title, "values": table_data}
                    extracted_tables.append(extracted_table)

        return extracted_tables

    except Exception as e:
        return []


def is_numerical(value: str) -> bool:
    """
    The function `is_numerical` checks if a given value is numerical, allowing for optional commas and a
    percentage sign at the end.

    Args:
      value (str): The value parameter is a string that represents a numerical value.

    Returns:
      The function is_numerical is returning a boolean value.
    """
    numerical_pattern = re.compile(r"^-?(\d{1,3}(,\d{3})*|\d+)?(\.\d+)?%?$")
    return bool(numerical_pattern.match(str(value).replace(",", "")))


def tables_to_html(list_of_tables: list, url: str) -> str:
    """
    The function `tables_to_html` takes a list of tables and a URL as input, and returns an HTML string
    containing the tables formatted with headers, rows, and a title, along with a clickable URL.

    Args:
      list_of_tables (list): A list of dictionaries, where each dictionary represents a table. Each
    dictionary should have two keys: "title" and "values". "title" represents the title of the table,
    and "values" represents the data in the table.
      url (str): The `url` parameter is a string that represents the URL of the table. It is used to
    create a hyperlink to the table in the generated HTML output.

    Returns:
      a string that contains HTML code representing tables.
    """
    try:
        html_tables = []

        for table in list_of_tables:
            if "title" in table and "values" in table:
                title = table["title"]
                values = table["values"]

                # Create the table header row
                headers = list(values[0].values())
                header_row = (
                    "<tr><th style='font-size: 16px; border: 1px solid black; padding: 10px; text-align: center;'>"
                    + "</th><th style='border: 1px solid black; padding: 10px; text-align: center;'>".join(
                        headers
                    )
                    + "</th></tr>"
                )
                html_table = [header_row]

                # Create rows with data
                for i in range(1, len(values)):
                    row = values[i]
                    row_values = list(row.values())
                    row_str = "<tr>"
                    for val in row_values:
                        # Check if the value is numerical
                        if is_numerical(str(val)):
                            align_style = "text-align: right;"
                        else:
                            align_style = "text-align: left;"
                        row_str += f"<td style='font-size: 14px; border: 1px solid black; padding: 10px; {align_style}'>{str(val)}</td>"
                    row_str += "</tr>"
                    html_table.append(row_str)

                # Create the table HTML structure
                table_html = (
                    "<table style='border-collapse: collapse;'><tbody>"
                    + "\n".join(html_table)
                    + "</tbody></table>"
                )

                # Create the title with proper heading
                title_html = (
                    f"<p style='font-weight: bold; font-size: 18px;'>{title}</p>"
                )

                # Add the url of the table
                url_html = f"<br><a href='{url}' style='float: right; font-size: 12px;'>{url}</a>"

                # Combine title and table HTML
                final_table_html = title_html + table_html + url_html
                html_tables.append(final_table_html)

        return "<br><br><br><br><br>".join(html_tables)

    except Exception as e:
        return ""


def add_table_to_doc(document: Document, table_title: str, table_values: list, url: str):
    try:
        document_clone = Document()
        document_clone.add_heading(table_title, level=2)
        table = document_clone.add_table(rows=1, cols=len(table_values[0]))
        table.style = "Table Grid"  # Applying table grid style to have borders

        # Set header rows as bold and add borders
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(table_values[0].keys()):
            cell = hdr_cells[i]
            cell.text = table_values[0][key]
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            for border in cell._element.xpath(".//*"):
                border_val = border.attrib.get("val")
                if border_val is not None:
                    border.attrib.clear()
                    border.attrib["val"] = border_val
                else:
                    border.attrib["val"] = "single"

        for row_data in table_values[1:]:
            row = table.add_row().cells
            for i, key in enumerate(row_data.keys()):
                cell = row[i]
                cell.text = row_data[key]
                if is_numerical(row_data[key]):
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                for border in cell._element.xpath(".//*"):
                    border_val = border.attrib.get("val")
                    if border_val is not None:
                        border.attrib.clear()
                        border.attrib["val"] = border_val
                    else:
                        border.attrib["val"] = "single"

        # Adding the table URL after the table with right alignment and as a hyperlink
        p = document_clone.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        add_hyperlink(p, url, url)

        # Append the modified content to the original document
        for element in document_clone.element.body:
            document.element.body.append(element)

    except Exception as e:
        # Revert changes on any exception
        pass
