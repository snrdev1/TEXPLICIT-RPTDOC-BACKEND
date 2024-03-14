import os
from datetime import datetime
from io import BytesIO, StringIO

import docx
import requests
from bson import ObjectId
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from google.cloud import storage
from pptx.util import Inches, Pt
from serpapi import GoogleSearch

from app.config import Config
from app.models.mongoClient import MongoClient
from .myDocumentsService import MyDocumentsService
from app.utils.llm.llm_summarizer import summarize_text
from app.utils.production import Production
from app.utils.web_scraper import WebScraper

from ..utils.document import add_hyperlink


class NewsService:
    @staticmethod
    def get_search_data(query, engine=0, count=10, start=0):
        """
        The get_search_data function takes in a query string, an engine number (0 for Google, 1 for Bing, 2 for Google Scholar), and the number of results to return.
        It then calls the appropriate API function based on the engine parameter and returns a list of dictionaries containing information about each result.

        Args:
            query: Search for a specific topic
            engine: Determine which api to use
            count: Specify the number of results to return

        Returns:
            A list of dictionaries
            :param start:
        """
        # Bing
        if engine == 1:
            results = NewsService._bing_api(query, count, start)
        # Google Scholar
        elif engine == 2:
            results = NewsService._google_scholar_api(query, count, start)
        # Google (default)
        else:
            results = NewsService._google_api(query, count, start)

        return results

    @staticmethod
    def get_news(articles, count=10):
        """
        The get_news function takes in a list of articles and returns the top 10 news articles.
            Args:
                articles (list): A list of dictionaries containing article information.

        Args:
            articles: Pass the list of articles from the rss feed
            count: Limit the number of articles to be processed

        Returns:
            A generator
        """
        article_count = 0
        print("Number of results : ", len(articles))
        for article in articles:
            print("New article processing started!!! : ", (article_count + 1))
            url = article["link"]
            result = WebScraper.scrape(url)

            if result:
                article_count += 1
                yield result

            if article_count == count:
                break

    @staticmethod
    def save_news_as_document(user_id, news, folder=None):
        """
        The function `save_news_as_document` saves news data as a document in a specified folder or the
        root directory, and returns the ID of the inserted document.

        Args:
          user_id: The user ID is a unique identifier for the user who is saving the news document. It
        is used to determine the root directory where the document will be saved.
          news: The "news" parameter is a dictionary that contains information about the news article.
        It typically includes the title and summary of the article.
          folder: The `folder` parameter is an optional parameter that represents the parent folder
        where the news document will be saved. It is a dictionary with two keys:

        Returns:
          the `inserted_id` of the saved news document.
        """

        # If parent folder is passed in then save root as the folder path else save document to root directory
        root_userid_substr = f"/{user_id}/"
        # print("Folder : ", folder)
        if folder:
            if folder == "/":
                root = root_userid_substr
            else:
                print("Folder root : ", folder["root"])
                if folder["root"] == root_userid_substr:
                    root = f"{folder['root']}{folder['originalFileName']}"
                else:
                    root = f"{folder['root']}/{folder['originalFileName']}"
        else:
            root = root_userid_substr
        print("Root : ", root)
        title = news.get("title", "")
        title = title + ".docx"
        content = news.get("summary", "")
        # print("Root : ", root)
        # Save news data to DB
        inserted_id = NewsService._save_doc_to_db(user_id, title, title, content, root)
        file_extension = "docx"
        if inserted_id:
            # Update the document virtual filename
            virtual_filename = inserted_id + "." + file_extension
            update_response = MyDocumentsService().update_virtual_filename(
                inserted_id, file_extension
            )

            # Create news document and save it to assets
            if folder == "/":
                NewsService._save_document(news, virtual_filename, user_id, "/")
            else:
                NewsService._save_document(news, virtual_filename, user_id, root)

        return inserted_id

    @staticmethod
    def _google_scholar_api(query, count=10, start=0):
        """
        The _google_scholar_api function takes a query and returns the top 10 results from Google Scholar.
            Args:
                query (str): The search term to be queried on Google Scholar.

        Args:
            query: Search for the query in Google Scholar
            count: Limit the number of results returned by the api

        Returns:
            A list of dictionaries
        """
        params = {
            "q": query,
            "api_key": Config.SERPAPI_KEY,
            "no_cache": Config.SERPAPI_NO_CACHE,
            "num": count,
            "start": start,
            "safe": "active",
            "engine": "google_scholar",
            # "tbm": 'nws'
        }

        client = GoogleSearch(params)
        results = client.get_dict()
        news_results = results["organic_results"]
        return news_results

    @staticmethod
    def _google_api(query, count=10, start=0):
        """
        The _google_api function takes in a query and count, and returns the top 10 results from Google.
            Args:
                query (str): The search term to be queried on Google.
                count (int): The number of results to return from the API call.

        Args:
            query: Search for the query that is passed in
            count: Determine how many results to return

        Returns:
            A list of dictionaries
        """
        params = {
            "q": query,
            "api_key": Config.SERPAPI_KEY,
            "no_cache": Config.SERPAPI_NO_CACHE,
            "num": count,
            "start": start,
            "safe": "active",
            "engine": "google",
            # "tbm": 'nws'
        }

        client = GoogleSearch(params)
        results = client.get_dict()
        print("Pagination results : ", results["serpapi_pagination"])
        news_results = results["organic_results"]

        return news_results

    @staticmethod
    def _bing_api(query, count=10, start=0):
        """
        The _bing_api function takes in a query and count, then returns the top 10 results from Bing.
            Args:
                query (str): The search term to be queried.
                count (int): The number of results to return. Defaults to 10 if not specified by user.

        Args:
            query: Search for a specific query
            count: Specify the number of results to return

        Returns:
            A list of dictionaries
        """
        params = {
            "q": query,
            "api_key": Config.SERPAPI_KEY,
            "no_cache": Config.SERPAPI_NO_CACHE,
            "count": count,
            "start": start,
            "safe": "active",
            "engine": "bing",
            # "tbm": 'nws'
        }

        client = GoogleSearch(params)
        results = client.get_dict()
        news_results = results["organic_results"]

        return news_results

    @staticmethod
    def _save_document(news, virtual_filename, user_id, folder=None):
        """
        The `_save_document` function saves a document in the DOCX format with the provided news
        information, virtual filename, and user ID.

        Args:
          news: The `news` parameter is a dictionary that contains information about a news article. It
        includes the following keys:
          virtual_filename: The virtual_filename parameter is a string that represents the name of the
        file to be saved. It is a virtual filename because it does not necessarily correspond to an
        actual file on the local file system. It is used to generate a unique filename for the document.
          user_id: The `user_id` parameter is the unique identifier for the user who is saving the
        document. It is used to create a folder specific to the user in which the document will be
        saved.

        Returns:
          None.
        """
        if Config.GCP_PROD_ENV:
            bucket = Production.get_users_bucket()
            
        try:
            title = news.get("title", "")
            summary = news.get("summary", "")
            # image_url = news.get('image', '')
            link = news.get("url", "")
            authors = ", ".join(news.get("authors", []))

            # Create a new document
            document = Document()

            # Add a heading for the item
            h = document.add_heading()
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = h.add_run(title)
            font = run.font
            font.size = Pt(16)

            # Add author names if present
            if len(authors) > 0:
                p_authors = document.add_paragraph()
                p_authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p_authors.add_run(authors)
                font = run.font
                font.size = Pt(10)

            document.add_paragraph()

            # # If there is an image url, add an image
            # if image_url:
            #     # Download the image from the URL
            #     image_from_url = urllib.request.urlopen(image_url)
            #     io_url = BytesIO()
            #     io_url.write(image_from_url.read())
            #     io_url.seek(0)
            #     # ================================================================= Not needed
            #     # response = requests.get(image_url)
            #     # print("Response: " , response)
            #     # image_content = BytesIO(response.content)
            #     # print("Image content: " , image_content)
            #     # print("doc : ", document)
            #     # document.add_picture(image_content, width=Inches(2), height=Inches(2))
            #     # ==================================================================
            #     try:
            #         document.add_picture(io_url, width=Inches(4), height=Inches(4))
            #     except Exception as e:
            #         print("Error:", e)

            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # print(summary)
            # Add a clickable link at the end of the summary
            if link:
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                add_hyperlink(p, link, link)

            # Add a paragraph for the summary
            p = document.add_paragraph()
            # print("Summary : ", summary)
            summarized_text = summarize_text(summary)
            # print(summarized_text)
            # run = p.add_run(summary)
            run = p.add_run(summarized_text)
            font = run.font
            font.size = Pt(12)

            timestamp = datetime.utcnow().timestamp()
            # print(f"SAVING FILE at {timestamp}!")
            # Ensure that the user upload folder exists
            # print("Folder in line 388 : ", folder)
            if folder == "/":
                user_folder_path = os.path.join(Config.USER_FOLDER, str(user_id))
            else:
                user_folder_path = os.path.join(Config.USER_FOLDER, folder[1:])
            os.makedirs(user_folder_path, exist_ok=True)
            # print("User folder path : ", user_folder_path)
            if folder == "/":
                file_save_path = MyDocumentsService.get_file_save_path(
                    virtual_filename, str(user_id), folder
                )
            else:
                new_folder_value = folder.replace(f"/{str(user_id)}", "")
                file_save_path = MyDocumentsService.get_file_save_path(
                    virtual_filename, str(user_id), new_folder_value
                )
            # print("File save path : ", file_save_path)
            document.save(file_save_path)
            print("Document saved to assets successfully!")
            if Config.GCP_PROD_ENV:
                if folder == "/":
                    destination_blob_name = str(user_id) + "/" + virtual_filename
                else:
                    destination_blob_name = folder[1:] + "/" + virtual_filename
                # print("Destination blob : ", destination_blob_name)
                blob = bucket.blob(destination_blob_name)
                # print("file_save_path : ", file_save_path)
                blob.upload_from_filename(file_save_path)
                print(
                    f"Uploaded news file {destination_blob_name} to bucket successfully!"
                )

        except Exception as e:
            print("Exception : ", e)
            return None

    @staticmethod
    def _save_doc_to_db(user_id, filename, title, content, root):
        """
        The function `_save_doc_to_db` saves a document to a database, ensuring that the filename is
        unique by appending an index if necessary.

        Args:
          user_id: The user ID of the user who is saving the document to the database.
          filename: The `filename` parameter is the original name of the file that you want to save to
        the database.
          title: The title of the document.
          content: The "content" parameter in the above code refers to the description or content of the
        document that you want to save to the database. It can be a string containing the text or any
        other relevant information about the document.
          root: The "root" parameter in the function represents the root directory or folder where the
        document will be saved. It is used to organize and categorize the documents in the database.

        Returns:
          the inserted ID of the document in the database if the insertion is successful. If the
        insertion fails, it returns None.
        """
        m_db = MongoClient.connect()
        print("Root : ", root)
        # Check if the file with the same name already exists in the collection
        existing_files = m_db[Config.MONGO_DOCUMENT_MASTER_COLLECTION].find(
            {"root": root}
        )

        # Create a set of existing filenames for efficient lookup
        existing_filenames = set(file["originalFileName"] for file in existing_files)

        # If the filename already exists, increment the index count
        index = 1
        new_filename = filename
        while new_filename in existing_filenames:
            # new_filename = f"{filename}({index})"
            # last_index = filename.rfind(".")
            file_name_without_extension, file_extension = os.path.splitext(filename)
            # print("File name without extension:", file_name_without_extension)
            new_filename = f"{file_name_without_extension}({index}){file_extension}"
            index = index + 1

        file_data = {
            "title": title,
            "description": content,
            "itemizedSummary": "",  # update when itemized summary of this record is generated
            "highlightsSummary": "",  # update when highlight summary of this record is generated
            "originalFileName": new_filename,  # Use the unique filename
            "virtualFileName": "",
            "createdBy": {"_id": ObjectId(user_id), "ref": "user"},
            "createdOn": datetime.utcnow(),
            "root": root,
            "type": "File",
            "embeddings": None,
            "usersWithAccess": [],
        }
        response = m_db[Config.MONGO_DOCUMENT_MASTER_COLLECTION].insert_one(file_data)

        if response:
            return str(response.inserted_id)

        return None
